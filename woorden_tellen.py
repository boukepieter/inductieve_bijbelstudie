import sqlite3
import sys
import docx
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_BREAK
import RGBColors
from random import randint
import pdb
import re
import urllib.request, urllib.parse, urllib.error
from html.parser import HTMLParser

# Parameters
titel = '1 Thes HSV'
threshold = 3
boek = "1TH"
hoofdstukken = range(1,6)
url = "http://herzienestatenvertaling.nl/copy.php?boek={:s}&hoofdstuk={:d}"
boeken = {"Efeziers" : "EPH", "Galaten" : "GAL", "Mattheus" : "MAT", "Genesis": "GEN", "1 Thessalonicenzen": "1TH"}
colors = RGBColors.COLORS_SMALL
noOfColors = 11
ignore = RGBColors.WORDS_IGNORE #[]

# Delen: [downloaden, tellen, wegschrijven]
execute = [1, 1, 1]


# create a subclass and override the handler methods
class MyHTMLParser(HTMLParser):
    write_data = False
    text = ""
    def handle_starttag(self, tag, attrs):
        return None
    def handle_endtag(self, tag):
        # pdb.set_trace()
        if tag == 'br':
            self.text = self.text + "\n"
    def handle_data(self, data):
        # pdb.set_trace()
        self.text = self.text + data

# scrape HSV from internet
if execute[0]:
    fh = open('{:s}.txt'.format(titel), 'w', encoding="UTF-8")
    for hoofdstuk in hoofdstukken:
        try:
            response = urllib.request.urlopen(url.format(boek, hoofdstuk))
        except urllib.error.URLError as e:
            print("failed")
        html = response.readlines()
        parser = MyHTMLParser()
        parser.feed(html[11].decode("UTF-8"))
        fh.write(parser.text)
    fh.close()

# set up database for counting
conn = sqlite3.connect('{:s}.sqlite'.format(titel))
cur = conn.cursor()
if execute[1]:
    cur.execute('''
    DROP TABLE IF EXISTS Counts''')
    cur.execute('''
    CREATE TABLE Counts (word TEXT, count INTEGER, color INTEGER)''')

    fh = open('{:s}.txt'.format(titel), encoding="UTF-8")
    for line in fh:
        if line[0].isalpha() : continue
        line = re.sub(r'[^\w\s]','',line)
        words = line.split()
        for word in words:
            word = word.lower()
            cur.execute('SELECT count FROM Counts WHERE word = ? ', (word, ))
            row = cur.fetchone()
            if word in ignore or word.isdigit():
                continue
            elif row is None:
                color = randint(0,noOfColors - 1)
                cur.execute('INSERT INTO Counts (word, count, color) VALUES ( ?, 1 , {:d})'.format(color), ( word, ))
            else:
                cur.execute('UPDATE Counts SET count=count+1 WHERE word = ?',
                    (word, ))
        # This statement commits outstanding changes to disk each
        # time through the loop - the program can be made faster
        # by moving the commit so it runs only after the loop completes
    conn.commit()

if execute[2]:
    doc = docx.Document()
    doc.add_heading(titel,1)
    par = doc.add_paragraph("")
    fh.close()
    fh = open('{:s}.txt'.format(titel), encoding="UTF-8")
    for line in fh:
        if line[0].isalpha() : continue
        words = line.split()
        for word in words:
            word_filt = re.sub(r'[^\w\s]','',word)
            run = par.add_run(word + " ")
            if word_filt.lower() in ignore or word.isdigit() or word_filt == "":
                continue
            cur.execute('SELECT count, color FROM Counts WHERE word = ? ', (word_filt.lower(), ))
            res = cur.fetchone()
            font = run.font
            if res[0] < threshold:
                font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            else :
                font.color.rgb = colors[res[1]]
                # font.highlight_color = res[1]

    conn.commit()
    run.add_break(WD_BREAK.PAGE)

    sqlstr = 'SELECT word, count, color FROM Counts WHERE count >= {:d} ORDER BY count DESC'.format(threshold)

    cur.execute(sqlstr)
    response = cur.fetchall()
    table = doc.add_table(cols=2, rows=len(response))
    for i in range(0,len(response)):
        cell = table.cell(i, 0)
        par = cell.paragraphs[0]
        run = par.add_run(response[i][0])
        font = run.font
        font.color.rgb = colors[response[i][2]]
        font.size =  Pt(15)
        # cell.text = response[i][0]
        cell = table.cell(i,1)
        par = cell.paragraphs[0]
        run = par.add_run(str(response[i][1]))
        font = run.font
        font.color.rgb = colors[response[i][2]]
        font.size = Pt(15)
    doc.save('{:s}_{:d}.docx'.format(titel, threshold))

cur.close()
